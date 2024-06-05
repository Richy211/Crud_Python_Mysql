from docx import Document

# Crear un nuevo documento de Word
doc = Document()

# Agregar título y texto al documento
doc.add_heading('Explicación del Código', level=1)

explanation = """
Claro, desglosaremos el código línea por línea:

### Imports

```python
import tkinter as tk
from tkinter import *
from tkinter import ttk
from tkinter import messagebox 
from Clientes import *
from Conexion import *
```

1. **`import tkinter as tk`**: Importa la biblioteca `tkinter` con el alias `tk`, la cual se utiliza para crear interfaces gráficas.
2. **`from tkinter import *`**: Importa todos los componentes del módulo `tkinter`, lo cual no es una práctica recomendada porque puede causar conflictos de nombres.
3. **`from tkinter import ttk`**: Importa el módulo `ttk` de `tkinter`, que proporciona widgets temáticos.
4. **`from tkinter import messagebox`**: Importa el módulo `messagebox` de `tkinter`, utilizado para mostrar cuadros de mensajes.
5. **`from Clientes import *` y `from Conexion import *`**: Importa todos los elementos de los módulos `Clientes` y `Conexion`, que probablemente contienen funciones o clases para manejar clientes y la conexión a la base de datos.

### Definición de la Clase

```python
class formularioClientes:
```

Define la clase `formularioClientes`, que contendrá toda la lógica para la interfaz gráfica de gestión de clientes.

### Método Constructor

```python
def __init__(self):
    self.base = None
    self.textBoxId = None
    self.textBoxNombres = None 
    self.textBoxApellidos = None 
    self.combo = None
    self.groupBox = None 
    self.tree = None 
    self.Formulario()
```

1. **`def __init__(self):`**: Define el constructor de la clase.
2. **Inicializa Variables**: Inicializa varias variables de instancia a `None`, que serán utilizadas para almacenar los widgets de la interfaz.
3. **`self.Formulario()`**: Llama al método `Formulario` para configurar la interfaz gráfica.

### Método para Guardar Registros

```python
def guardarRegistros(self):
    try:
        # Verify if the widgets are initialized 
        if self.textBoxNombres is None or self.textBoxApellidos is None or self.combo is None: 
            print("Los widget no estan inicializados")
            return 
        nombres = self.textBoxNombres.get()
        apellidos = self.textBoxApellidos.get()
        sexo = self.combo.get() 
        
        CClientes.ingresarClientes(nombres, apellidos, sexo) 
        messagebox.showinfo("Informacion", "Los datos han sido guardados")
        
        # Clear the fields
        self.textBoxNombres.delete(0, END)
        self.textBoxApellidos.delete(0, END)
        
    except ValueError as error:
        print("Error al ingresar los datos {}".format(error))
```

1. **`def guardarRegistros(self):`**: Define el método para guardar registros.
2. **Verificación de Widgets**: Verifica si los widgets de entrada están inicializados.
3. **Obtención de Valores**: Obtiene los valores de los campos de texto y del combo box.
4. **Guardar en Base de Datos**: Llama a `ingresarClientes` del módulo `CClientes` para guardar los datos en la base de datos.
5. **Mensaje de Confirmación**: Muestra un mensaje de confirmación.
6. **Limpiar Campos**: Limpia los campos de texto.
7. **Manejo de Errores**: Captura y muestra cualquier error que ocurra durante la operación.

### Método para Configurar la Interfaz Gráfica

```python
def Formulario(self):  
    try:
        self.base = Tk()
        self.base.geometry("1200x300")
        self.base.title("Formulario Python")
        
        self.groupBox = LabelFrame(self.base, text="Datos del Personal", padx=5, pady=5)
        self.groupBox.grid(row=0, column=0, padx=10, pady=10)
        
        labelId = Label(self.groupBox, text="Id:", width=13, font=("arial", 12)).grid(row=0, column=0)
        self.textBoxId = Entry(self.groupBox)
        self.textBoxId.grid(row=0, column=1)
        
        labelNombre = Label(self.groupBox, text="Nombre:", width=13, font=("arial", 12)).grid(row=1, column=0)
        self.textBoxNombres = Entry(self.groupBox)
        self.textBoxNombres.grid(row=1, column=1)
        
        labelApellido = Label(self.groupBox, text="Apellido:", width=13, font=("arial", 12)).grid(row=2, column=0)
        self.textBoxApellidos = Entry(self.groupBox)
        self.textBoxApellidos.grid(row=2, column=1)
        
        labelSexo = Label(self.groupBox, text="Sexo:", width=13, font=("arial", 12)).grid(row=3, column=0)
        seleccionSexo = tk.StringVar()
        self.combo = ttk.Combobox(self.groupBox, values=["Masculino", "Femenino"], textvariable=seleccionSexo)
        self.combo.grid(row=3, column=1)
        seleccionSexo.set("Masculino")
        
        Button(self.groupBox, text="Guardar", width=10, command=self.guardarRegistros).grid(row=4, column=0)
        Button(self.groupBox, text="Modificar", width=10).grid(row=4, column=1)
        Button(self.groupBox, text="Eliminar", width=10).grid(row=4, column=2)
        
        self.groupBox = LabelFrame(self.base, text="Lista del Personal", padx=5, pady=5)
        self.groupBox.grid(row=0, column=1, padx=5, pady=5)
        # Create a TreeView 
        
        # Configure the columns
        
        self.tree = ttk.Treeview(self.groupBox, columns=("Id", "Nombres", "Apellidos", "Sexo"), show='headings', height=5)
        self.tree.column("# 1", anchor=CENTER)
        self.tree.heading("# 1", text="ID")
        self.tree.column("# 2", anchor=CENTER)
        self.tree.heading("# 2", text="Nombres")
        self.tree.column("# 3", anchor=CENTER)
        self.tree.heading("# 3", text="Apellidos")
        self.tree.column("# 4", anchor=CENTER)
        self.tree.heading("# 4", text="Sexo")
        
        self.tree.pack()
        
        self.base.mainloop()
    
    except ValueError as error:
        print("Error al mostrar la interfaz, error: {}".format(error))
```

1. **`def Formulario(self):`**: Define el método para configurar la interfaz gráfica.
2. **Crear Ventana Principal**: Crea la ventana principal de la aplicación.
   - `self.base = Tk()`: Inicializa la ventana principal.
   - `self.base.geometry("1200x300")`: Define el tamaño de la ventana.
   - `self.base.title("Formulario Python")`: Establece el título de la ventana.
3. **Crear `LabelFrame` para Datos Personales**:
   - `self.groupBox = LabelFrame(self.base, text="Datos del Personal", padx=5, pady=5)`: Crea un `LabelFrame` para agrupar los campos de entrada de datos personales.
   - `self.groupBox.grid(row=0, column=0, padx=10, pady=10)`: Coloca el `LabelFrame` en la ventana.
4. **Agregar Campos de Entrada**:
   - Crea y coloca etiquetas (`Label`) y campos de entrada (`Entry`) para ID, Nombre, Apellido y Sexo.
   - Para el sexo, utiliza un `ttk.Combobox` con opciones "Masculino" y "Femenino".
5. **Botones de Acción**:
   - Crea botones para "Guardar", "Modificar" y "Eliminar".
   - El botón "Guardar" está vinculado al método `guardarRegistros`.
6. **Crear `LabelFrame` para Lista de Personal**:
   - Crea otro `LabelFrame` para la lista de personal.
   - Dentro de este `LabelFrame`, crea un `Treeview` para mostrar los datos en formato de tabla.
   - Configura las columnas del `Treeview` (ID, Nombres, Apellidos, Sexo).
7. **Iniciar el Bucle Principal**:
   - `self.base.mainloop()`: Inicia el bucle principal de la interfaz gráfica, permitiendo la interacción del usuario.
8. **Manejo de Errores**: Captura y muestra cualquier error que ocurra durante la configuración de la interfaz gráfica.

### Ejecución Principal

```python
formularioClientes()
```

- Crea una instancia de la clase `formularioClientes`, lo que automáticamente llama al constructor y, por ende, inicializa la interfaz gráfica.

En resumen, este código crea una interfaz gráfica para gestionar información de clientes, permitiendo al usuario ingresar datos, guardarlos en una base de datos y mostrarlos en una tabla dentro de la misma interfaz.

"""

doc.add_paragraph(explanation)

# Guardar el documento
doc.save('ExplicacionCodigo.docx')
